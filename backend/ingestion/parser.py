"""Document parsing module for PDF, DOCX, and HTML files.

Extracts raw text with page-level granularity from legal documents
in the data directory hierarchy.
"""

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import pdfplumber
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".html", ".htm"}


@dataclass
class ParsedPage:
    """A single page of extracted text."""
    page_number: int  # 1-indexed
    text: str


@dataclass
class ParsedDocument:
    """A fully parsed document with page-level text and source metadata."""
    file_path: str
    file_name: str
    subfolder: str  # e.g., "statues", "case_law", "training"
    pages: list[ParsedPage] = field(default_factory=list)
    full_text: str = ""
    total_pages: int = 0


def parse_pdf(file_path: str) -> ParsedDocument:
    """Parse a PDF file using pdfplumber, extracting text page by page."""
    path = Path(file_path)
    pages: list[ParsedPage] = []

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            pages.append(ParsedPage(page_number=i + 1, text=text))

    full_text = "\n\n".join(p.text for p in pages)

    return ParsedDocument(
        file_path=str(path.resolve()),
        file_name=path.name,
        subfolder="",  # set by caller
        pages=pages,
        full_text=full_text,
        total_pages=len(pages),
    )


def parse_docx(file_path: str) -> ParsedDocument:
    """Parse a DOCX file using python-docx, extracting paragraph text."""
    path = Path(file_path)
    doc = DocxDocument(file_path)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    # DOCX has no native page concept; treat as single page
    pages = [ParsedPage(page_number=1, text=full_text)] if full_text else []

    return ParsedDocument(
        file_path=str(path.resolve()),
        file_name=path.name,
        subfolder="",
        pages=pages,
        full_text=full_text,
        total_pages=1,
    )


def parse_html(file_path: str) -> ParsedDocument:
    """Parse an HTML file using BeautifulSoup, preserving header structure."""
    path = Path(file_path)

    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, "html.parser")

    # Remove script and style elements
    for tag in soup(["script", "style"]):
        tag.decompose()

    full_text = soup.get_text(separator="\n")

    # Treat as single page
    pages = [ParsedPage(page_number=1, text=full_text)] if full_text.strip() else []

    return ParsedDocument(
        file_path=str(path.resolve()),
        file_name=path.name,
        subfolder="",
        pages=pages,
        full_text=full_text,
        total_pages=1,
    )


def parse_file(file_path: str) -> ParsedDocument:
    """Dispatch to the appropriate parser based on file extension."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext in (".html", ".htm"):
        return parse_html(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext} ({file_path})")


def _resolve_subfolder(file_path: Path, data_dir: Path) -> str:
    """Determine the subfolder name relative to the data directory.

    For data/statues/statute_1.pdf -> "statues"
    For data/case_law/2023AP001664.pdf -> "case_law"
    """
    try:
        relative = file_path.relative_to(data_dir)
        parts = relative.parts
        return parts[0] if len(parts) > 1 else ""
    except ValueError:
        return ""


def parse_directory(data_dir: str) -> list[ParsedDocument]:
    """Walk the data directory recursively and parse all supported files.

    Args:
        data_dir: Path to the root data directory containing subfolders.

    Returns:
        List of ParsedDocument objects, one per successfully parsed file.
    """
    data_path = Path(data_dir).resolve()
    if not data_path.is_dir():
        raise FileNotFoundError(f"Data directory not found: {data_dir}")

    files = sorted(
        f for f in data_path.rglob("*")
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    )

    if not files:
        logger.warning(f"No supported files found in {data_dir}")
        return []

    documents: list[ParsedDocument] = []

    for i, file_path in enumerate(files, 1):
        logger.info(f"Parsing {file_path.name} ({i}/{len(files)})")
        try:
            doc = parse_file(str(file_path))
            doc.subfolder = _resolve_subfolder(file_path, data_path)
            documents.append(doc)
        except Exception as e:
            logger.error(f"Failed to parse {file_path}: {e}")
            continue

    logger.info(f"Successfully parsed {len(documents)}/{len(files)} files")
    return documents
